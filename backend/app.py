import asyncio
import base64
import json
import logging
import math
import os
import random
import sqlite3
import subprocess
import uuid
from datetime import datetime, timedelta
from typing import Any, Dict, List, Optional

import httpx
from fastapi import FastAPI, HTTPException, Request, File, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from passlib.context import CryptContext
from pydantic import BaseModel, Field
from pathlib import Path

from prisma import Prisma
from prisma.errors import DataError

from backend.services.predictive_lab import predictive_lab
from backend.services.social_impact import build_social_impact

# Import ML endpoints
try:
    from backend.ml.inference.ml_endpoints import router as ml_router
    ML_AVAILABLE = True
except ImportError:
    ML_AVAILABLE = False
    print("Warning: ML endpoints not available. Run: python backend/ml/models/all_models.py")

AI_MODEL_DEFAULT = "gemini-1.5-flash"
AI_API_KEY = (
    os.getenv("GENAI_API_KEY")
    or os.getenv("GOOGLE_GENAI_API_KEY")
    or os.getenv("GEMINI_API_KEY")
)

PROJECT_ROOT = Path(__file__).resolve().parents[1]

ALLOWED_ORIGINS = os.getenv(
    "ALLOWED_ORIGINS",
    "http://localhost:5173,http://localhost:3000,http://localhost:3002,http://localhost:3004,http://localhost:3005",
)
ALLOW_ALL_ORIGINS = os.getenv("ALLOW_ALL_ORIGINS", "false").lower() in {"1", "true", "yes"}
ALLOW_ORIGIN_REGEX = os.getenv("ALLOWED_ORIGIN_REGEX", r"http://localhost:\d+")

cors_origins = [origin.strip() for origin in ALLOWED_ORIGINS.split(",") if origin.strip()]
if ALLOW_ALL_ORIGINS:
    cors_origins = ["*"]
    allow_credentials = False
    allow_origin_regex = None
else:
    allow_credentials = True
    allow_origin_regex = ALLOW_ORIGIN_REGEX or None

app = FastAPI()
app.add_middleware(
    CORSMiddleware,
    allow_origins=cors_origins,
    allow_credentials=allow_credentials,
    allow_methods=["*"],
    allow_headers=["*"],
    allow_origin_regex=allow_origin_regex,
)

# Include ML router if available
if ML_AVAILABLE:
    app.include_router(ml_router)
    print("✓ ML endpoints registered successfully!")

prisma = Prisma()
pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")


def _resolve_sqlite_path(url: str) -> str:
    if url.startswith("file:"):
        url = url[5:]
    if url.startswith("//"):
        url = url[2:]
    if url.startswith("./") or url.startswith(".\\"):
        url = url[2:]
        url = os.path.join(PROJECT_ROOT, url)
    path = url.split("?", 1)[0]
    if os.name == "nt" and len(path) >= 3 and path[0] in ("/", "\\") and path[2] == ":":
        path = path[1:]
    return os.path.abspath(path)

_SQLITE_DB_PATH = _resolve_sqlite_path(os.getenv("DATABASE_URL", "file:./data/databases/dev.db"))


@app.on_event("startup")
async def on_startup():
    await prisma.connect()


@app.on_event("shutdown")
async def on_shutdown():
    await prisma.disconnect()


def _uuid(value: Optional[str] = None) -> str:
    return value or str(uuid.uuid4())


def _iso(dt: Optional[datetime]) -> Optional[str]:
    return dt.isoformat() if dt else None


def _safe_json_loads(payload: Optional[str]) -> Dict[str, Any]:
    if not payload:
        return {}
    try:
        data = json.loads(payload)
        if isinstance(data, dict):
            return data
        return {"value": data}
    except json.JSONDecodeError:
        return {"content": payload}


def _hash_password(raw: Optional[str] = None) -> str:
    secret = raw or os.getenv("DEFAULT_USER_PASSWORD", "Synapse@2025")
    return pwd_context.hash(secret)

# --- Criptografia em repouso (Fernet) ---
_DATA_KEY = os.getenv("DATA_ENCRYPTION_KEY")
try:
    from cryptography.fernet import Fernet
    _fernet = Fernet(_DATA_KEY) if _DATA_KEY else None
except Exception:
    _fernet = None

def _enc(value: Optional[str]) -> Optional[str]:
    if not value:
        return value
    if not _fernet:
        return value
    return _fernet.encrypt(value.encode("utf-8")).decode("utf-8")

def _dec(value: Optional[str]) -> Optional[str]:
    if not value:
        return value
    if not _fernet:
        return value
    try:
        return _fernet.decrypt(value.encode("utf-8")).decode("utf-8")
    except Exception:
        return value

# --- Rate Limiting (IP/rota) ---
# Aumentado para 200 requisições por minuto para suportar carregamento inicial
RATE_MAX = int(os.getenv("RATE_MAX", "200"))
RATE_WINDOW_SEC = int(os.getenv("RATE_WINDOW_SEC", "60"))
_RATE_BUCKETS: Dict[str, List[float]] = {}

@app.middleware("http")
async def rate_limit(request: Request, call_next):
    ip = request.client.host if request.client else "unknown"
    key = f"{ip}:{request.url.path}"
    now = datetime.utcnow().timestamp()
    window_start = now - RATE_WINDOW_SEC
    bucket = _RATE_BUCKETS.get(key, [])
    bucket = [t for t in bucket if t >= window_start]
    if len(bucket) >= RATE_MAX:
        retry = int(bucket[0] + RATE_WINDOW_SEC - now)
        from fastapi.responses import PlainTextResponse
        return PlainTextResponse("Too Many Requests", status_code=429, headers={"Retry-After": str(max(retry, 1))})
    bucket.append(now)
    _RATE_BUCKETS[key] = bucket
    return await call_next(request)

# --- Auditoria ---
async def audit(db: Prisma, user_id: Optional[str], action: str, details: str):
    try:
        await db.execute_raw(
            "INSERT INTO logs_auditoria (id, idUsuario, acao, detalhes, dataHora) VALUES (?, ?, ?, ?, ?)",
            (str(uuid.uuid4()), user_id or "system", action, details, datetime.utcnow().isoformat()),
        )
    except Exception:
        pass


def _normalize_activity_type(t: str) -> str:
    if not t:
        return "RESUMO"
    t = t.upper().replace(" ", "_").replace("-", "_")
    
    mapping = {
        "SUMMARY": "RESUMO",
        "RESUMO": "RESUMO",
        "FLASHCARDS": "FLASHCARD",
        "FLASHCARD": "FLASHCARD",
        "QUIZ": "QUIZ",
        "SIMULATION": "SIMULADO",
        "SIMULADO": "SIMULADO",
        "CASE_STUDY": "ESTUDO_DE_CASO",
        "ESTUDO_DE_CASO": "ESTUDO_DE_CASO",
        "CASE": "ESTUDO_DE_CASO",
        "MIND_MAP": "MAPA_MENTAL",
        "MAP_MIND": "MAPA_MENTAL",
        "MAPA_MENTAL": "MAPA_MENTAL",
        "CLOZE": "CLOZE",
        "CLOZE_TEST": "CLOZE",
        "COMPARISON": "COMPARACAO",
        "COMPARACAO": "COMPARACAO",
        "COMPARISON_TABLE": "COMPARACAO",
        "SELF_EXPLANATION": "AUTO_EXPLICACAO",
        "AUTO_EXPLICACAO": "AUTO_EXPLICACAO"
    }
    
    return mapping.get(t, t)

def _module_from_activity(activity) -> Dict[str, Any]:
    module = _safe_json_loads(activity.conteudo)
    module.setdefault("id", activity.id)
    
    raw_type = module.get("tipo") or activity.tipo
    module["type"] = _normalize_activity_type(raw_type)
    
    module.setdefault("title", module.get("titulo") or module.get("title") or "Atividade")
    module.setdefault("description", module.get("descricao") or module.get("description") or "")
    module.setdefault("isCompleted", bool(module.get("isCompleted", False)))
    module["xpReward"] = int(module.get("xpReward") or activity.xpReward or 10)
    module["estimatedTimeMin"] = int(module.get("estimatedTimeMin") or module.get("tempoEstimadoMin") or activity.tempoEstimadoMin or 5)
    module["avaliativa"] = bool(module.get("avaliativa") or activity.avaliativa)
    return module


def _course_status(progress: int) -> str:
    if progress >= 100:
        return "CONCLUIDO"
    if progress > 0:
        return "EM_ANDAMENTO"
    return "NAO_INICIADO"


def map_course(record) -> Dict[str, Any]:
    atividades = sorted(record.atividades or [], key=lambda item: item.ordem)
    modules = [_module_from_activity(item) for item in atividades]
    completed = sum(1 for module in modules if module.get("isCompleted"))
    progress = round((completed / len(modules)) * 100) if modules else 0
    total_xp = sum(int(module.get("xpReward") or 0) for module in modules)

    tag_names = []
    for link in record.tags or []:
        tag_entity = getattr(link, "tag", None)
        if tag_entity and tag_entity.nome:
            tag_names.append(tag_entity.nome)

    return {
        "id": record.id,
        "title": record.titulo,
        "description": record.descricao,
        "category": record.categoria or (record.tipoCurso or "").title(),
        "content": record.textoExtraido,
        "modules": modules,
        "progress": progress,
        "totalXP": total_xp,
        "thumbnailUrl": record.thumbnailUrl,
        "createdAt": _iso(record.criadoEm),
        "tags": tag_names,
        "isRecommended": (record.tipoCurso or "").upper() == "RECOMENDADO",
        "status": _course_status(progress),
    }


def map_user(record) -> Dict[str, Any]:
    matriculas = record.matriculas or []
    courses_assigned = len(matriculas)
    courses_completed = sum(1 for item in matriculas if (item.status or "").upper() == "CONCLUIDO")
    courses_late = sum(
        1
        for item in matriculas
        if (item.status or "").upper() in {"ATRASADO", "REPROVADO_SIMULADO"}
    )
    avatar = record.avatarUrl or f"https://i.pravatar.cc/150?u={record.id}"

    return {
        "id": record.id,
        "name": record.nome,
        "email": _dec(record.email),
        "cpf": _dec(record.cpf),
        "role": record.cargo or "Colaborador",
        "teamId": record.idEquipe,
        "avatarUrl": avatar,
        "totalXP": record.totalXp or 0,
        "level": record.nivel or 1,
        "streakDays": record.diasSequencia or 0,
        "coursesAssigned": courses_assigned,
        "coursesCompleted": courses_completed,
        "coursesLate": courses_late,
    }


def map_team(record) -> Dict[str, Any]:
    usuarios = record.usuarios or []
    member_count = len(usuarios)
    progresses = [
        matricula.progresso or 0
        for user in usuarios
        for matricula in (user.matriculas or [])
    ]
    avg_completion = round(sum(progresses) / len(progresses)) if progresses else 0
    scores = [
        matricula.notaFinal or 0
        for user in usuarios
        for matricula in (user.matriculas or [])
        if matricula.notaFinal is not None
    ]
    avg_score = round(sum(scores) / len(scores)) if scores else 0

    return {
        "id": record.id,
        "name": record.nome,
        "area": record.descricao or "",  # Legacy: kept for backwards compatibility
        "areaId": record.idArea,
        "areaName": record.area.nome if record.area else None,
        "managerId": record.idGestor,
        "stats": {
            "memberCount": member_count,
            "avgCompletionRate": avg_completion,
            "avgSimulationScore": avg_score,
        },
    }


def map_role(record) -> Dict[str, Any]:
    return {
        "id": record.id,
        "name": record.nome,
        "teamId": record.idEquipe,
        "description": record.descricao,
    }


def map_area(record) -> Dict[str, Any]:
    return {
        "id": record.id,
        "name": record.nome,
        "description": record.descricao,
    }


def map_enrollment(record) -> Dict[str, Any]:
    return {
        "id": record.id,
        "courseId": record.idCurso,
        "collaboratorId": record.idUsuario,
        "isRequired": bool(record.ehObrigatorio),
        "assignedAt": _iso(record.atribuidoEm),
        "dueDate": _iso(record.prazo),
        "status": record.status,
        "progress": record.progresso or 0,
        "finalScore": record.notaFinal,
        "lastAccessAt": _iso(record.ultimoAcesso),
    }


def _calculate_team_stats_from_sqlite(conn: sqlite3.Connection, team_id: str) -> Dict[str, int]:
    member_rows = conn.execute("SELECT id FROM usuarios WHERE idEquipe = ?", (team_id,)).fetchall()
    member_ids = [row["id"] for row in member_rows]
    member_count = len(member_ids)
    avg_completion = 0
    avg_score = 0

    if member_ids:
        placeholders = ",".join(["?"] * len(member_ids))
        enrollment_rows = conn.execute(
            f"SELECT progresso, notaFinal FROM matriculas WHERE idUsuario IN ({placeholders})",
            tuple(member_ids),
        ).fetchall()
        completions = [
            row["progresso"] if row["progresso"] is not None else 0
            for row in enrollment_rows
        ]
        if completions:
            avg_completion = round(sum(completions) / len(completions))

        scores = [row["notaFinal"] for row in enrollment_rows if row["notaFinal"] is not None]
        if scores:
            avg_score = round(sum(scores) / len(scores))

    return {
        "memberCount": member_count,
        "avgCompletionRate": avg_completion,
        "avgSimulationScore": avg_score,
    }


def _load_teams_from_sqlite() -> List[Dict[str, Any]]:
    if not os.path.exists(_SQLITE_DB_PATH):
        return []

    conn = sqlite3.connect(_SQLITE_DB_PATH)
    conn.row_factory = sqlite3.Row
    try:
        rows = conn.execute(
            "SELECT id, nome, descricao, idGestor FROM equipes"
        ).fetchall()
        fallback: List[Dict[str, Any]] = []
        for row in rows:
            fallback.append(
                {
                    "id": row["id"],
                    "name": row["nome"],
                    "area": row["descricao"] or "",
                    "managerId": row["idGestor"],
                    "stats": _calculate_team_stats_from_sqlite(conn, row["id"]),
                }
            )
        return fallback
    finally:
        conn.close()


class CoursePayload(BaseModel):
    id: Optional[str] = None
    title: str
    description: Optional[str] = None
    category: Optional[str] = None
    content: Optional[str] = None
    tags: List[str] = Field(default_factory=list)
    modules: List[Dict[str, Any]] = Field(default_factory=list)
    thumbnailUrl: Optional[str] = None
    tipoCurso: Optional[str] = None
    tipoArquivo: Optional[str] = None
    urlArquivo: Optional[str] = None
    statusProcessamento: Optional[str] = None


class CourseUpdatePayload(BaseModel):
    title: Optional[str] = None
    description: Optional[str] = None
    category: Optional[str] = None
    content: Optional[str] = None
    tags: Optional[List[str]] = None
    modules: Optional[List[Dict[str, Any]]] = None
    thumbnailUrl: Optional[str] = None
    tipoCurso: Optional[str] = None
    tipoArquivo: Optional[str] = None
    urlArquivo: Optional[str] = None
    statusProcessamento: Optional[str] = None


class UserPayload(BaseModel):
    id: Optional[str] = None
    name: str
    email: str
    cpf: Optional[str] = None
    role: Optional[str] = None
    teamId: Optional[str] = None
    avatarUrl: Optional[str] = None
    totalXP: int = 0
    level: int = 1
    streakDays: int = 0
    aceitouTermos: bool = False
    preferenciaAcessibilidade: Optional[str] = None


class UserUpdatePayload(BaseModel):
    name: Optional[str] = None
    email: Optional[str] = None
    cpf: Optional[str] = None
    role: Optional[str] = None
    teamId: Optional[str] = None
    avatarUrl: Optional[str] = None
    totalXP: Optional[int] = None
    level: Optional[int] = None
    streakDays: Optional[int] = None
    aceitouTermos: Optional[bool] = None
    preferenciaAcessibilidade: Optional[str] = None


class TeamPayload(BaseModel):
    name: str
    area: Optional[str] = None
    managerId: Optional[str] = None


class TeamUpdatePayload(BaseModel):
    name: Optional[str] = None
    area: Optional[str] = None
    managerId: Optional[str] = None


class RolePayload(BaseModel):
    id: Optional[str] = None
    name: str
    teamId: Optional[str] = None
    description: Optional[str] = None


class RoleUpdatePayload(BaseModel):
    name: Optional[str] = None
    teamId: Optional[str] = None
    description: Optional[str] = None


class AreaPayload(BaseModel):
    name: str
    description: Optional[str] = None


class AreaUpdatePayload(BaseModel):
    name: Optional[str] = None
    description: Optional[str] = None


class EnrollmentPayload(BaseModel):
    id: Optional[str] = None
    courseId: str
    collaboratorId: str
    isRequired: bool = False
    status: str = "NAO_INICIADO"
    progress: int = 0
    finalScore: Optional[float] = None
    assignedAt: Optional[str] = None
    dueDate: Optional[str] = None
    lastAccessAt: Optional[str] = None


class AIPayload(BaseModel):
    prompt: str
    model: Optional[str] = None
    responseSchema: Optional[Dict[str, Any]] = None
    responseMimeType: Optional[str] = None
    config: Optional[Dict[str, Any]] = None
    systemInstruction: Optional[str] = None


class BioData(BaseModel):
    userId: str
    bpm: int
    gsr: int
    movement: int


class BioPredictPayload(BaseModel):
    userId: Optional[str] = None


class SessionPayload(BaseModel):
    id: Optional[str] = None
    userId: str
    activityId: str
    checkInBioId: Optional[str] = None
    modoRecomendado: Optional[str] = None


class SessionCompletePayload(BaseModel):
    score: Optional[float] = None
    completedAt: Optional[str] = None


class InteractionPayload(BaseModel):
    id: Optional[str] = None
    userId: str
    sessionId: Optional[str] = None
    activityId: str
    isCorrect: Optional[bool] = None
    responseTimeMs: int
    attempts: int = 1


class ReviewPayload(BaseModel):
    id: Optional[str] = None
    userId: str
    activityId: str
    easinessFactor: float = 2.5
    interval: int = 0
    repetitions: int = 0
    lastReview: Optional[str] = None
    nextReview: Optional[str] = None


async def resolve_cargo_id(name: Optional[str]) -> Optional[str]:
    if not name:
        return None
    existing = await prisma.cargo.find_unique(where={"nome": name})
    if existing:
        return existing.id
    created = await prisma.cargo.create(data={"nome": name})
    return created.id


async def ensure_tag_id(tag_name: str) -> str:
    value = tag_name.strip()
    if not value:
        raise HTTPException(status_code=400, detail="Tag inválida")
    existing = await prisma.tag.find_unique(where={"nome": value})
    if existing:
        return existing.id
    created = await prisma.tag.create(data={"nome": value})
    return created.id


async def sync_course_tags(course_id: str, tags: List[str]) -> None:
    await prisma.materialtag.delete_many(where={"idMaterial": course_id})
    for tag_name in tags:
        tag_id = await ensure_tag_id(tag_name)
        await prisma.materialtag.create(
            data={
                "idMaterial": course_id,
                "idTag": tag_id,
            }
        )


async def sync_course_modules(course_id: str, modules: List[Dict[str, Any]]) -> None:
    existing = await prisma.atividadeaprendizado.find_many(
        where={"idMaterialFonte": course_id}
    )
    existing_map = {item.id: item for item in existing}
    keep_ids: set[str] = set()

    for order, raw in enumerate(modules):
        module_id = _uuid(raw.get("id"))
        keep_ids.add(module_id)
        payload = {
            "tipo": raw.get("type") or raw.get("tipo") or "RESUMO",
            "conteudo": json.dumps(raw, ensure_ascii=False),
            "xpReward": int(raw.get("xpReward") or raw.get("xp_reward") or 10),
            "tempoEstimadoMin": int(raw.get("estimatedTimeMin") or raw.get("tempoEstimadoMin") or 5),
            "ordem": order,
            "avaliativa": bool(
                raw.get("avaliativa")
                or str(raw.get("type") or "").upper() in {"SIMULADO", "QUIZ"}
            ),
        }
        if module_id in existing_map:
            await prisma.atividadeaprendizado.update(
                where={"id": module_id},
                data=payload,
            )
        else:
            await prisma.atividadeaprendizado.create(
                data={
                    "id": module_id,
                    "idMaterialFonte": course_id,
                    **payload,
                }
            )

    for activity in existing:
        if activity.id not in keep_ids:
            await prisma.atividadeaprendizado.delete(where={"id": activity.id})


async def fetch_course_record(course_id: str):
    record = await prisma.materialfonte.find_unique(
        where={"id": course_id},
        include={
            "atividades": True,
            "tags": {"include": {"tag": True}},
        },
    )
    if not record:
        raise HTTPException(status_code=404, detail="Curso não encontrado")
    return record


async def ensure_exists(fetcher, where: Dict[str, Any], message: str):
    record = await fetcher(where=where)
    if not record:
        raise HTTPException(status_code=404, detail=message)
    return record


def _parse_datetime(value: Optional[str], fallback: Optional[datetime] = None) -> datetime:
    if value:
        try:
            return datetime.fromisoformat(value.replace("Z", "+00:00"))
        except ValueError:
            pass
    return fallback or datetime.utcnow()


@app.get("/courses")
async def list_courses():
    records = await prisma.materialfonte.find_many(
        include={
            "atividades": True,
            "tags": {"include": {"tag": True}},
        },
    )
    ordered = sorted(records, key=lambda record: record.criadoEm or datetime.min, reverse=True)
    return [map_course(record) for record in ordered]


@app.get("/courses/{course_id}")
async def get_course(course_id: str):
    record = await fetch_course_record(course_id)
    return map_course(record)


@app.post("/courses", status_code=201)
async def create_course(payload: CoursePayload):
    course_id = _uuid(payload.id)
    tipo_curso = (payload.tipoCurso or "RECOMENDADO").upper()
    tipo_arquivo = payload.tipoArquivo or "AI_GENERATED"
    url_arquivo = payload.urlArquivo or f"https://cdn.synapse/{course_id}.pdf"
    status = payload.statusProcessamento or "CONCLUIDO"

    await prisma.materialfonte.create(
        data={
            "id": course_id,
            "titulo": payload.title,
            "descricao": payload.description,
            "categoria": payload.category,
            "tipoCurso": tipo_curso,
            "thumbnailUrl": payload.thumbnailUrl,
            "urlArquivo": url_arquivo,
            "tipoArquivo": tipo_arquivo,
            "statusProcessamento": status,
            "textoExtraido": payload.content,
        }
    )

    if payload.tags:
        await sync_course_tags(course_id, payload.tags)
    if payload.modules:
        await sync_course_modules(course_id, payload.modules)

    record = await fetch_course_record(course_id)
    return map_course(record)


@app.put("/courses/{course_id}")
async def update_course(course_id: str, payload: CourseUpdatePayload):
    update_data: Dict[str, Any] = {}
    if payload.title is not None:
        update_data["titulo"] = payload.title
    if payload.description is not None:
        update_data["descricao"] = payload.description
    if payload.category is not None:
        update_data["categoria"] = payload.category
    if payload.content is not None:
        update_data["textoExtraido"] = payload.content
    if payload.thumbnailUrl is not None:
        update_data["thumbnailUrl"] = payload.thumbnailUrl
    if payload.tipoCurso is not None:
        update_data["tipoCurso"] = payload.tipoCurso.upper()
    if payload.tipoArquivo is not None:
        update_data["tipoArquivo"] = payload.tipoArquivo
    if payload.urlArquivo is not None:
        update_data["urlArquivo"] = payload.urlArquivo
    if payload.statusProcessamento is not None:
        update_data["statusProcessamento"] = payload.statusProcessamento

    if update_data:
        await prisma.materialfonte.update(where={"id": course_id}, data=update_data)

    if payload.tags is not None:
        await sync_course_tags(course_id, payload.tags)
    if payload.modules is not None:
        await sync_course_modules(course_id, payload.modules)

    record = await fetch_course_record(course_id)
    return map_course(record)


@app.delete("/courses/{course_id}")
async def delete_course(course_id: str):
    await fetch_course_record(course_id)
    await prisma.materialfonte.delete(where={"id": course_id})
    return {"deleted": True}


@app.get("/users")
async def list_users():
    records = await prisma.usuario.find_many(include={"matriculas": True})
    ordered = sorted(records, key=lambda record: record.nome or "")
    return [map_user(record) for record in ordered]


@app.post("/users", status_code=201)
async def create_user(payload: UserPayload):
    user_id = _uuid(payload.id)
    cpf = payload.cpf or f"{random.randint(10**10, 10**11 - 1):011d}"
    cargo_id = await resolve_cargo_id(payload.role)

    enc_email = _enc(payload.email)
    enc_cpf = _enc(cpf)
    await prisma.usuario.create(
        data={
            "id": user_id,
            "nome": payload.name,
            "email": enc_email,
            "cpf": enc_cpf,
            "hashSenha": _hash_password(),
            "avatarUrl": payload.avatarUrl,
            "papel": "COLABORADOR",
            "idEquipe": payload.teamId,
            "idCargo": cargo_id,
            "cargo": payload.role,
            "totalXp": payload.totalXP,
            "nivel": payload.level,
            "diasSequencia": payload.streakDays,
            "aceitouTermos": payload.aceitouTermos,
            "preferenciaAcessibilidade": payload.preferenciaAcessibilidade,
        }
    )

    record = await prisma.usuario.find_unique(
        where={"id": user_id},
        include={"matriculas": True},
    )
    return map_user(record)


@app.put("/users/{user_id}")
async def update_user(user_id: str, payload: UserUpdatePayload):
    await ensure_exists(prisma.usuario.find_unique, {"id": user_id}, "Usuário não encontrado")
    cargo_id = await resolve_cargo_id(payload.role)
    update_data: Dict[str, Any] = {}

    if payload.name is not None:
        update_data["nome"] = payload.name
    if payload.email is not None:
        update_data["email"] = _enc(payload.email)
    if payload.cpf is not None:
        update_data["cpf"] = _enc(payload.cpf)
    if payload.teamId is not None:
        update_data["idEquipe"] = payload.teamId
    if payload.avatarUrl is not None:
        update_data["avatarUrl"] = payload.avatarUrl
    if payload.totalXP is not None:
        update_data["totalXp"] = payload.totalXP
    if payload.level is not None:
        update_data["nivel"] = payload.level
    if payload.streakDays is not None:
        update_data["diasSequencia"] = payload.streakDays
    if payload.role is not None:
        update_data["cargo"] = payload.role
    if cargo_id is not None:
        update_data["idCargo"] = cargo_id
    if payload.aceitouTermos is not None:
        update_data["aceitouTermos"] = payload.aceitouTermos
    if payload.preferenciaAcessibilidade is not None:
        update_data["preferenciaAcessibilidade"] = payload.preferenciaAcessibilidade

    if update_data:
        await prisma.usuario.update(where={"id": user_id}, data=update_data)

    record = await prisma.usuario.find_unique(
        where={"id": user_id},
        include={"matriculas": True},
    )
    return map_user(record)


@app.delete("/users/{user_id}")
async def delete_user(user_id: str):
    await ensure_exists(prisma.usuario.find_unique, {"id": user_id}, "Usuário não encontrado")
    await prisma.usuario.delete(where={"id": user_id})
    return {"deleted": True}


@app.get("/teams")
async def list_teams():
    try:
        records = await prisma.equipe.find_many(
            include={"usuarios": {"include": {"matriculas": True}}, "area": True},
        )
        return [map_team(record) for record in records]
    except DataError as exc:
        logging.warning("Falling back to SQLite while loading teams: %s", exc)
        return await asyncio.to_thread(_load_teams_from_sqlite)


@app.get("/api/leaderboard")
async def leaderboard(scope: str = "team"):
    if scope != "team":
        # futuro: suporte a global/user
        scope = "team"
    teams = await prisma.equipe.find_many(include={"usuarios": True})
    ranking = []
    for t in teams:
        total_xp = sum((u.totalXp or 0) for u in (t.usuarios or []))
        avg_level = (
            sum((u.nivel or 0) for u in (t.usuarios or [])) / max(1, len(t.usuarios or []))
        )
        ranking.append({"teamId": t.id, "teamName": t.nome, "totalXP": total_xp, "avgLevel": int(avg_level)})
    ranking.sort(key=lambda r: r["totalXP"], reverse=True)
    return {"scope": scope, "ranking": ranking}


@app.get("/api/missions/current")
async def missions_current():
    # Missões estáticas simples para a semana
    start = datetime.utcnow() - timedelta(days=datetime.utcnow().weekday())
    end = start + timedelta(days=7)
    missions = [
        {"id": "m-review-3", "title": "Concluir 3 revisões (SR)", "rewardXP": 500},
        {"id": "m-perfect-sim", "title": "Tirar 100 em um simulado", "rewardXP": 800},
        {"id": "m-collab-help", "title": "Ajudar 1 colega (feedback)", "rewardXP": 300},
    ]
    return {"weekStart": start.isoformat(), "weekEnd": end.isoformat(), "missions": missions}


@app.post("/api/missions/complete")
async def missions_complete(userId: str, missionId: str):
    await prisma.execute_raw(
        "CREATE TABLE IF NOT EXISTS mission_progress (id TEXT PRIMARY KEY, userId TEXT, missionId TEXT, doneAt TEXT)"
    )
    await prisma.execute_raw(
        "INSERT INTO mission_progress (id, userId, missionId, doneAt) VALUES (?, ?, ?, ?)",
        (str(uuid.uuid4()), userId, missionId, datetime.utcnow().isoformat()),
    )
    return {"ok": True}


@app.get("/api/season")
async def season_status():
    # Temporada baseada em XP acumulado
    users = await prisma.usuario.find_many()
    total_xp = sum((u.totalXp or 0) for u in users)
    milestones = [5000, 15000, 30000, 50000]
    reached = sum(1 for m in milestones if total_xp >= m)
    return {"totalXP": total_xp, "milestones": milestones, "reached": reached}


@app.post("/teams", status_code=201)
async def create_team(payload: TeamPayload):
    record = await prisma.equipe.create(
        data={
            "id": _uuid(),
            "nome": payload.name,
            "descricao": payload.area,  # Legacy field
            "idGestor": payload.managerId,
        },
        include={"usuarios": {"include": {"matriculas": True}}, "area": True},
    )
    return map_team(record)


@app.put("/teams/{team_id}")
async def update_team(team_id: str, payload: TeamUpdatePayload):
    await ensure_exists(prisma.equipe.find_unique, {"id": team_id}, "Equipe não encontrada")
    update_data = {}
    if payload.name is not None:
        update_data["nome"] = payload.name
    if payload.area is not None:
        update_data["descricao"] = payload.area
    if payload.managerId is not None:
        update_data["idGestor"] = payload.managerId
    if update_data:
        await prisma.equipe.update(where={"id": team_id}, data=update_data)

    record = await prisma.equipe.find_unique(
        where={"id": team_id},
        include={"usuarios": {"include": {"matriculas": True}}, "area": True},
    )
    return map_team(record)


@app.delete("/teams/{team_id}")
async def delete_team(team_id: str):
    await ensure_exists(prisma.equipe.find_unique, {"id": team_id}, "Equipe não encontrada")
    await prisma.equipe.delete(where={"id": team_id})
    return {"deleted": True}


@app.get("/areas")
async def list_areas():
    records = await prisma.area.find_many(
        include={"equipes": {"include": {"usuarios": True}}},
    )
    return [map_area(record) for record in records]


@app.post("/areas", status_code=201)
async def create_area(payload: AreaPayload):
    record = await prisma.area.create(
        data={
            "id": _uuid(),
            "nome": payload.name,
            "descricao": payload.description,
        },
        include={"equipes": {"include": {"usuarios": True}}},
    )
    return map_area(record)


@app.put("/areas/{area_id}")
async def update_area(area_id: str, payload: AreaUpdatePayload):
    await ensure_exists(prisma.area.find_unique, {"id": area_id}, "Área não encontrada")
    update_data = {}
    if payload.name is not None:
        update_data["nome"] = payload.name
    if payload.description is not None:
        update_data["descricao"] = payload.description
    if update_data:
        await prisma.area.update(where={"id": area_id}, data=update_data)

    record = await prisma.area.find_unique(
        where={"id": area_id},
        include={"equipes": {"include": {"usuarios": True}}},
    )
    return map_area(record)


@app.delete("/areas/{area_id}")
async def delete_area(area_id: str):
    await ensure_exists(prisma.area.find_unique, {"id": area_id}, "Área não encontrada")
    # Check if area has teams
    teams_count = await prisma.equipe.count(where={"idArea": area_id})
    if teams_count > 0:
        raise HTTPException(status_code=400, detail="Área possui times associados")
    await prisma.area.delete(where={"id": area_id})
    return {"deleted": True}


@app.get("/roles")
async def list_roles(teamId: Optional[str] = None):
    where_clause = {}
    if teamId:
        where_clause["idEquipe"] = teamId
    records = await prisma.cargo.find_many(where=where_clause, include={"equipe": True})
    ordered = sorted(records, key=lambda record: record.nome or "")
    return [map_role(record) for record in ordered]


@app.post("/roles", status_code=201)
async def create_role(payload: RolePayload):
    data = {
        "id": _uuid(payload.id),
        "nome": payload.name,
    }
    if payload.teamId is not None:
        data["idEquipe"] = payload.teamId
    if payload.description is not None:
        data["descricao"] = payload.description

    record = await prisma.cargo.create(data=data, include={"equipe": True})
    return map_role(record)


@app.put("/roles/{role_id}")
async def update_role(role_id: str, payload: RoleUpdatePayload):
    await ensure_exists(prisma.cargo.find_unique, {"id": role_id}, "Cargo não encontrado")
    update_data = {}
    if payload.name is not None:
        update_data["nome"] = payload.name
    if payload.teamId is not None:
        update_data["idEquipe"] = payload.teamId
    if payload.description is not None:
        update_data["descricao"] = payload.description
    if update_data:
        await prisma.cargo.update(where={"id": role_id}, data=update_data)

    record = await prisma.cargo.find_unique(
        where={"id": role_id},
        include={"equipe": True},
    )
    return map_role(record)


@app.delete("/roles/{role_id}")
async def delete_role(role_id: str):
    role = await prisma.cargo.find_unique(where={"id": role_id})
    if not role:
        raise HTTPException(status_code=404, detail="Cargo não encontrado")
    usage = await prisma.usuario.count(where={"idCargo": role_id})
    if usage > 0:
        raise HTTPException(status_code=400, detail="Cargo em uso por colaboradores")
    await prisma.cargo.delete(where={"id": role_id})
    return {"deleted": True}


@app.get("/enrollments")
async def list_enrollments():
    records = await prisma.matricula.find_many()
    return [map_enrollment(record) for record in records]


@app.post("/enrollments", status_code=201)
async def create_enrollment(payload: EnrollmentPayload):
    await ensure_exists(prisma.usuario.find_unique, {"id": payload.collaboratorId}, "Usuário não encontrado")
    await ensure_exists(prisma.materialfonte.find_unique, {"id": payload.courseId}, "Curso não encontrado")
    record = await prisma.matricula.create(
        data={
            "id": _uuid(payload.id),
            "idUsuario": payload.collaboratorId,
            "idCurso": payload.courseId,
            "status": payload.status,
            "progresso": payload.progress,
            "notaFinal": payload.finalScore,
            "ehObrigatorio": payload.isRequired,
            "atribuidoEm": _parse_datetime(payload.assignedAt),
            "prazo": _parse_datetime(payload.dueDate, datetime.utcnow() + timedelta(days=15)),
            "ultimoAcesso": _parse_datetime(payload.lastAccessAt, datetime.utcnow()),
        }
    )
    return map_enrollment(record)


@app.put("/enrollments/{enrollment_id}")
async def update_enrollment(enrollment_id: str, payload: EnrollmentPayload):
    await ensure_exists(prisma.matricula.find_unique, {"id": enrollment_id}, "Matrícula não encontrada")
    data: Dict[str, Any] = {}
    if payload.status is not None:
        data["status"] = payload.status
    if payload.progress is not None:
        data["progresso"] = payload.progress
    if payload.finalScore is not None:
        data["notaFinal"] = payload.finalScore
    if payload.isRequired is not None:
        data["ehObrigatorio"] = payload.isRequired
    if payload.dueDate is not None:
        data["prazo"] = _parse_datetime(payload.dueDate)
    if payload.lastAccessAt is not None:
        data["ultimoAcesso"] = _parse_datetime(payload.lastAccessAt)
    if payload.assignedAt is not None:
        data["atribuidoEm"] = _parse_datetime(payload.assignedAt)
    if data:
        record = await prisma.matricula.update(where={"id": enrollment_id}, data=data)
    else:
        record = await prisma.matricula.find_unique(where={"id": enrollment_id})
    return map_enrollment(record)


@app.delete("/enrollments/{enrollment_id}")
async def delete_enrollment(enrollment_id: str):
    await ensure_exists(prisma.matricula.find_unique, {"id": enrollment_id}, "Matrícula não encontrada")
    await prisma.matricula.delete(where={"id": enrollment_id})
    return {"deleted": True}


@app.post("/api/ai")
async def generate_ai(payload: AIPayload):
    api_key = AI_API_KEY
    if not api_key:
        raise HTTPException(status_code=400, detail="Configure a chave da API Gemini/GénAI")

    model = payload.model or AI_MODEL_DEFAULT
    url = f"https://generativelanguage.googleapis.com/v1beta/models/{model}:generateContent"
    request_body: Dict[str, Any] = {
        "contents": [
            {
                "parts": [
                    {"text": payload.prompt},
                ]
            }
        ]
    }
    if payload.responseSchema:
        request_body["responseSchema"] = payload.responseSchema
    if payload.responseMimeType:
        request_body["responseMimeType"] = payload.responseMimeType
    if payload.config:
        request_body["generationConfig"] = payload.config
    if payload.systemInstruction:
        request_body["systemInstruction"] = {
            "parts": [{"text": payload.systemInstruction}]
        }

    async with httpx.AsyncClient(timeout=60) as client:
        response = await client.post(url, params={"key": api_key}, json=request_body)

    if response.status_code != 200:
        raise HTTPException(status_code=response.status_code, detail=response.text)

    data = response.json()
    text_output = ""
    candidates = data.get("candidates", [])
    if candidates:
        parts = candidates[0].get("content", {}).get("parts", [])
        text_output = "".join(part.get("text", "") for part in parts)

    return {"text": text_output, "raw": data}


@app.get("/api/ai/stream")
async def stream_ai(request: Request, prompt: str, model: Optional[str] = None):
    api_key = AI_API_KEY
    if not api_key:
        raise HTTPException(status_code=400, detail="Configure a chave da API Gemini/GénAI")
    from fastapi.responses import StreamingResponse
    mdl = model or AI_MODEL_DEFAULT
    url = f"https://generativelanguage.googleapis.com/v1beta/models/{mdl}:generateContent"
    body = {"contents": [{"parts": [{"text": prompt}]}]}

    async def event_stream():
        try:
            async with httpx.AsyncClient(timeout=60) as client:
                resp = await client.post(url, params={"key": api_key}, json=body)
            if resp.status_code != 200:
                yield f"event: error\ndata: {resp.text}\n\n"
                return
            data = resp.json()
            text = ""
            candidates = data.get("candidates", [])
            if candidates:
                parts = candidates[0].get("content", {}).get("parts", [])
                text = "".join(part.get("text", "") for part in parts)
            # emitir em chunks artificiais
            for i in range(0, len(text), 64):
                chunk = text[i:i+64]
                yield f"data: {json.dumps({'chunk': chunk})}\n\n"
            yield "event: end\ndata: {}\n\n"
        except Exception as e:
            yield f"event: error\ndata: {json.dumps({'error': str(e)})}\n\n"

    return StreamingResponse(event_stream(), media_type="text/event-stream")


@app.post("/api/extract-text")
async def extract_text_from_file(file: UploadFile = File(...)):
    """
    Extrai texto de arquivos PDF, DOCX, DOC, TXT e MD
    """
    try:
        # Obter extensão do arquivo
        file_extension = Path(file.filename).suffix.lower() if file.filename else ""

        # Ler o conteúdo do arquivo
        contents = await file.read()

        # Processar de acordo com o tipo
        if file_extension in ['.txt', '.md']:
            # Para arquivos de texto simples
            text = contents.decode('utf-8', errors='ignore')
            return {"text": text, "filename": file.filename}

        elif file_extension == '.pdf':
            # Extrair texto de PDF usando PyPDF2
            try:
                import PyPDF2
                import io

                pdf_file = io.BytesIO(contents)
                pdf_reader = PyPDF2.PdfReader(pdf_file)

                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"

                return {"text": text.strip(), "filename": file.filename}
            except Exception as e:
                raise HTTPException(
                    status_code=400,
                    detail=f"Erro ao extrair texto do PDF: {str(e)}"
                )

        elif file_extension in ['.docx', '.doc']:
            # Extrair texto de DOCX usando python-docx
            try:
                from docx import Document
                import io

                doc_file = io.BytesIO(contents)
                doc = Document(doc_file)

                text = ""
                for paragraph in doc.paragraphs:
                    text += paragraph.text + "\n"

                # Também extrair texto de tabelas
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            text += cell.text + " "
                        text += "\n"

                return {"text": text.strip(), "filename": file.filename}
            except Exception as e:
                raise HTTPException(
                    status_code=400,
                    detail=f"Erro ao extrair texto do DOCX: {str(e)}"
                )

        else:
            raise HTTPException(
                status_code=400,
                detail=f"Formato de arquivo não suportado: {file_extension}. Use .txt, .md, .pdf, .doc ou .docx"
            )

    except Exception as e:
        logging.error(f"Erro ao processar arquivo: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail=f"Erro ao processar arquivo: {str(e)}"
        )


@app.post("/api/iot/checkin")
async def receive_biofeedback(data: BioData):
    nivel_estresse = min(100, int((data.bpm / 120) * 50 + (data.gsr / 10) * 50))
    nivel_foco = max(0, 100 - nivel_estresse)

    user = await prisma.usuario.find_unique(where={"id": data.userId})
    if not user:
        await prisma.usuario.create(
            data={
                "id": data.userId,
                "nome": "Colaborador IoT",
                "email": f"{data.userId}@synapse.ai",
                "cpf": f"{random.randint(10**10, 10**11 - 1):011d}",
                "hashSenha": _hash_password(),
                "papel": "COLABORADOR",
                "cargo": "Colaborador",
            }
        )

    now = datetime.utcnow()
    checkin_id = str(uuid.uuid4())
    await prisma.checkinbio.create(
        data={
            "id": checkin_id,
            "idUsuario": data.userId,
            "horasSono": random.randint(5, 8),
            "qualidadeSono": random.randint(6, 10),
            "nivelFoco": nivel_foco,
            "nivelEstresse": nivel_estresse,
            "nivelFadiga": random.randint(10, 60),
            "origemDados": "SENSOR_REAL_TIME",
            "dadosBrutosSensor": json.dumps({"bpm": data.bpm, "gsr": data.gsr, "movement": data.movement}),
            "dataHora": now,
            "diaDaSemana": now.weekday(),
            "horaDoDia": now.hour,
        }
    )

    await prisma.logauditoria.create(
        data={
            "id": str(uuid.uuid4()),
            "idUsuario": data.userId,
            "acao": "CHECKIN_IOT",
            "detalhes": f"Recebido biofeedback ID {checkin_id}",
        }
    )

    recommendation_mode = "ALTA_PERFORMANCE"
    if nivel_estresse > 70:
        recommendation_mode = "PAUSA_GUIADA"
    elif nivel_foco > 80:
        recommendation_mode = "MICRO_APRENDIZADO"

    return {
        "status": "synced",
        "checkin_id": checkin_id,
        "bio_analysis": {"stress": nivel_estresse, "focus": nivel_foco, "mode": recommendation_mode},
    }


def _neuro_actions(projection: Dict[str, float]) -> List[Dict[str, Any]]:
    stress = projection.get("stress", 0)
    focus = projection.get("focus", 0)
    actions = []
    if stress >= 70:
        actions.append({"title": "Pausa guiada", "description": "Respiração 4-7-8 e playlist calmante."})
    if focus < 55:
        actions.append({"title": "Sprint Pomodoro Verde", "description": "15 minutos com protocolo de luz natural."})
    if not actions:
        actions.append({"title": "Microlearning Turbo", "description": "Liberar missão sustentável sugerida pelo bot."})
    return actions


@app.post("/api/iot/predict")
async def predict_bio_state(payload: BioPredictPayload):
    result = await predictive_lab.predict_for_user(prisma, payload.userId)
    projection = result.get("projection", {})
    response = {
        **result,
        "actions": _neuro_actions(projection),
    }
    await audit(prisma, payload.userId or "system", "IOT_PREDICT", f"stress={projection.get('stress', 0):.1f}")
    return response


@app.get("/api/dashboard/stats")
async def get_stats():
    aggregate = await prisma.checkinbio.aggregate(
        _avg={"nivelEstresse": True, "nivelFoco": True}
    )
    avg_data = getattr(aggregate, "_avg", {}) or {}
    return {
        "media_empresa_stress": avg_data.get("nivelEstresse") or 0,
        "media_empresa_foco": avg_data.get("nivelFoco") or 0,
    }


# --- ANALYTICS ---

R_ANALYTICS_SCRIPT = PROJECT_ROOT / "scripts" / "analytics" / "analise_cluster.R"
R_ANALYTICS_OUTPUT = R_ANALYTICS_SCRIPT.parent / "insight_r_plot.png"
R_ANALYTICS_SUMMARY = R_ANALYTICS_SCRIPT.parent / "insight_r_summary.json"


@app.get("/api/analytics/r/report")
async def analytics_r_report():
    if not R_ANALYTICS_SCRIPT.exists():
        raise HTTPException(status_code=404, detail="Relatório R indisponível.")

    try:
        subprocess.run(
            ["Rscript", str(R_ANALYTICS_SCRIPT)],
            cwd=str(R_ANALYTICS_SCRIPT.parent),
            check=True,
            capture_output=True,
        )
    except FileNotFoundError:
        raise HTTPException(
            status_code=500,
            detail="Rscript não encontrado no ambiente. Instale R/Rscript para gerar o relatório.",
        )
    except subprocess.CalledProcessError as exc:
        logging.error("Erro ao executar Rscript: %s", exc.stderr.decode("utf-8", errors="ignore"))
        raise HTTPException(status_code=500, detail="Falha ao gerar relatório em R.")

    if not R_ANALYTICS_OUTPUT.exists():
        raise HTTPException(status_code=500, detail="R não gerou arquivo de saída.")

    with open(R_ANALYTICS_OUTPUT, "rb") as img:
        encoded = base64.b64encode(img.read()).decode("ascii")

    summary_payload: Dict[str, Any] = {}
    if R_ANALYTICS_SUMMARY.exists():
        try:
            summary_payload = json.loads(R_ANALYTICS_SUMMARY.read_text(encoding="utf-8"))
        except Exception:
            summary_payload = {}

    return {"report": encoded, "summary": summary_payload}


@app.get("/api/analytics/overview")
async def analytics_overview():
    teams = await prisma.equipe.find_many(include={"usuarios": True})
    data = []
    for t in teams:
        user_ids = [u.id for u in (t.usuarios or [])]
        if not user_ids:
            data.append({"teamId": t.id, "teamName": t.nome, "avgStress": 0, "avgFocus": 0})
            continue
        checkins = await prisma.checkinbio.find_many(where={"idUsuario": {"in": user_ids}})
        if not checkins:
            data.append({"teamId": t.id, "teamName": t.nome, "avgStress": 0, "avgFocus": 0})
            continue
        stress_values = [c.nivelEstresse or 0 for c in checkins]
        focus_values = [c.nivelFoco or 0 for c in checkins]
        avg_stress = sum(stress_values) / len(stress_values) if stress_values else 0
        avg_focus = sum(focus_values) / len(focus_values) if focus_values else 0
        data.append({"teamId": t.id, "teamName": t.nome, "avgStress": avg_stress, "avgFocus": avg_focus})
    return {"teams": data}


# --- SESSÕES DE APRENDIZADO ---

@app.post("/api/sessions")
async def create_session(payload: SessionPayload):
    session_id = _uuid(payload.id)
    now = datetime.utcnow()

    session = await prisma.sessaoaprendizado.create(
        data={
            "id": session_id,
            "idUsuario": payload.userId,
            "idCheckInBio": payload.checkInBioId,
            "modoRecomendado": payload.modoRecomendado or "PADRAO",
            "iniciadoEm": now,
        }
    )

    await audit(prisma, payload.userId, "SESSION_START", f"Sessão {session_id} iniciada")

    return {
        "id": session.id,
        "userId": session.idUsuario,
        "checkInBioId": session.idCheckInBio,
        "modoRecomendado": session.modoRecomendado,
        "startedAt": _iso(session.iniciadoEm),
        "completedAt": _iso(session.concluidoEm),
    }


@app.put("/api/sessions/{session_id}/complete")
async def complete_session(session_id: str, payload: SessionCompletePayload):
    completed_at = datetime.fromisoformat(payload.completedAt) if payload.completedAt else datetime.utcnow()

    session = await prisma.sessaoaprendizado.update(
        where={"id": session_id},
        data={
            "concluidoEm": completed_at,
        }
    )

    await audit(prisma, session.idUsuario, "SESSION_COMPLETE", f"Sessão {session_id} concluída com score {payload.score}")

    return {
        "id": session.id,
        "userId": session.idUsuario,
        "completedAt": _iso(session.concluidoEm),
        "score": payload.score,
    }


@app.get("/api/users/{user_id}/sessions")
async def get_user_sessions(user_id: str, limit: int = 50):
    sessions = await prisma.sessaoaprendizado.find_many(
        where={"idUsuario": user_id},
        order_by={"iniciadoEm": "desc"},
        take=limit,
    )

    return {
        "sessions": [
            {
                "id": s.id,
                "userId": s.idUsuario,
                "checkInBioId": s.idCheckInBio,
                "modoRecomendado": s.modoRecomendado,
                "startedAt": _iso(s.iniciadoEm),
                "completedAt": _iso(s.concluidoEm),
            }
            for s in sessions
        ]
    }


# --- INTERAÇÕES ---

@app.post("/api/interactions")
async def create_interaction(payload: InteractionPayload):
    interaction_id = _uuid(payload.id)

    interaction = await prisma.interacao.create(
        data={
            "id": interaction_id,
            "idUsuario": payload.userId,
            "idSessao": payload.sessionId,
            "idAtividade": payload.activityId,
            "estaCorreto": payload.isCorrect,
            "tempoRespostaMs": payload.responseTimeMs,
            "tentativas": payload.attempts,
        }
    )

    return {
        "id": interaction.id,
        "userId": interaction.idUsuario,
        "sessionId": interaction.idSessao,
        "activityId": interaction.idAtividade,
        "isCorrect": interaction.estaCorreto,
        "responseTimeMs": interaction.tempoRespostaMs,
        "attempts": interaction.tentativas,
        "createdAt": _iso(interaction.criadoEm),
    }


@app.get("/api/activities/{activity_id}/interactions")
async def get_activity_interactions(activity_id: str, user_id: Optional[str] = None, limit: int = 100):
    where_clause = {"idAtividade": activity_id}
    if user_id:
        where_clause["idUsuario"] = user_id

    interactions = await prisma.interacao.find_many(
        where=where_clause,
        order_by={"criadoEm": "desc"},
        take=limit,
    )

    return {
        "interactions": [
            {
                "id": i.id,
                "userId": i.idUsuario,
                "sessionId": i.idSessao,
                "activityId": i.idAtividade,
                "isCorrect": i.estaCorreto,
                "responseTimeMs": i.tempoRespostaMs,
                "attempts": i.tentativas,
                "createdAt": _iso(i.criadoEm),
            }
            for i in interactions
        ]
    }


# --- SPACED REPETITION ---

@app.get("/api/users/{user_id}/reviews")
async def get_user_reviews(user_id: str, due_only: bool = True):
    where_clause = {"idUsuario": user_id}
    if due_only:
        where_clause["proximaRevisao"] = {"lte": datetime.utcnow()}

    reviews = await prisma.revisaosr.find_many(
        where=where_clause,
        order_by={"proximaRevisao": "asc"},
    )

    return {
        "reviews": [
            {
                "id": r.id,
                "userId": r.idUsuario,
                "activityId": r.idAtividade,
                "easinessFactor": r.easinessFactor,
                "interval": r.intervalo,
                "repetitions": r.repeticoes,
                "lastReview": _iso(r.ultimaRevisao),
                "nextReview": _iso(r.proximaRevisao),
            }
            for r in reviews
        ]
    }


@app.post("/api/reviews")
async def create_or_update_review(payload: ReviewPayload):
    existing = await prisma.revisaosr.find_first(
        where={
            "idUsuario": payload.userId,
            "idAtividade": payload.activityId,
        }
    )

    last_review = datetime.fromisoformat(payload.lastReview) if payload.lastReview else datetime.utcnow()
    next_review = datetime.fromisoformat(payload.nextReview) if payload.nextReview else (last_review + timedelta(days=payload.interval or 1))

    if existing:
        review = await prisma.revisaosr.update(
            where={"id": existing.id},
            data={
                "easinessFactor": payload.easinessFactor,
                "intervalo": payload.interval,
                "repeticoes": payload.repetitions,
                "ultimaRevisao": last_review,
                "proximaRevisao": next_review,
            }
        )
    else:
        review = await prisma.revisaosr.create(
            data={
                "id": _uuid(payload.id),
                "idUsuario": payload.userId,
                "idAtividade": payload.activityId,
                "easinessFactor": payload.easinessFactor,
                "intervalo": payload.interval,
                "repeticoes": payload.repetitions,
                "ultimaRevisao": last_review,
                "proximaRevisao": next_review,
            }
        )

    await audit(prisma, payload.userId, "REVIEW_UPDATE", f"Review atualizada para atividade {payload.activityId}")

    return {
        "id": review.id,
        "userId": review.idUsuario,
        "activityId": review.idAtividade,
        "easinessFactor": review.easinessFactor,
        "interval": review.intervalo,
        "repetitions": review.repeticoes,
        "lastReview": _iso(review.ultimaRevisao),
        "nextReview": _iso(review.proximaRevisao),
    }


@app.put("/api/reviews/{review_id}")
async def update_review(review_id: str, payload: ReviewPayload):
    last_review = datetime.fromisoformat(payload.lastReview) if payload.lastReview else datetime.utcnow()
    next_review = datetime.fromisoformat(payload.nextReview) if payload.nextReview else (last_review + timedelta(days=payload.interval or 1))

    review = await prisma.revisaosr.update(
        where={"id": review_id},
        data={
            "easinessFactor": payload.easinessFactor,
            "intervalo": payload.interval,
            "repeticoes": payload.repetitions,
            "ultimaRevisao": last_review,
            "proximaRevisao": next_review,
        }
    )

    return {
        "id": review.id,
        "userId": review.idUsuario,
        "activityId": review.idAtividade,
        "easinessFactor": review.easinessFactor,
        "interval": review.intervalo,
        "repetitions": review.repeticoes,
        "lastReview": _iso(review.ultimaRevisao),
        "nextReview": _iso(review.proximaRevisao),
    }



@app.get("/api/dashboard/collaborator/{user_id}")
async def dashboard_collaborator(user_id: str):
    '''
    Retorna dados agregados para o dashboard do colaborador
    '''
    # Buscar dados do usuário
    user = await prisma.usuario.find_unique(
        where={"id": user_id},
        include={"equipe": True}
    )

    if not user:
        raise HTTPException(status_code=404, detail="Usuário não encontrado")

    # Stats pessoais
    stats = {
        "xp": user.totalXp or 0,
        "nivel": user.nivel or 1,
        "streak": user.diasSequencia or 0,
        "avatarUrl": user.avatarUrl,
        "nome": user.nome,
    }

    # Calcular XP para próximo nível (fórmula: 500 * nivel^1.5)
    current_level = user.nivel or 1
    xp_for_next = int(500 * ((current_level + 1) ** 1.5))
    xp_current_level = int(500 * (current_level ** 1.5))
    xp_progress = (user.totalXp or 0) - xp_current_level
    xp_needed = xp_for_next - xp_current_level

    stats["xpForNextLevel"] = xp_needed
    stats["xpProgressInLevel"] = xp_progress
    stats["xpPercentageInLevel"] = int((xp_progress / xp_needed) * 100) if xp_needed > 0 else 100

    # Contagem de cursos por status
    matriculas = await prisma.matricula.find_many(
        where={"idUsuario": user_id},
        include={"curso": True}
    )

    cursos_stats = {
        "emAndamento": 0,
        "atrasados": 0,
        "concluidos": 0,
        "naoIniciados": 0,
        "total": len(matriculas),
        "progressoMedio": 0
    }

    progresso_total = 0
    for m in matriculas:
        if m.status == "EM_ANDAMENTO":
            cursos_stats["emAndamento"] += 1
        elif m.status == "ATRASADO":
            cursos_stats["atrasados"] += 1
        elif m.status == "CONCLUIDO":
            cursos_stats["concluidos"] += 1
        elif m.status == "NAO_INICIADO":
            cursos_stats["naoIniciados"] += 1
        progresso_total += (m.progresso or 0)

    if len(matriculas) > 0:
        cursos_stats["progressoMedio"] = int(progresso_total / len(matriculas))

    stats["cursos"] = cursos_stats

    # Ranking na equipe (se pertence a uma equipe)
    ranking = None
    if user.idEquipe:
        team_users_all = await prisma.usuario.find_many(
            where={"idEquipe": user.idEquipe}
        )
        # Sort by XP descending
        team_users = sorted(team_users_all, key=lambda u: u.totalXp or 0, reverse=True)

        ranking = {
            "posicao": 0,
            "total": len(team_users),
            "top3": [],
            "xpDoLider": 0,
            "xpParaProximo": 0
        }

        for idx, u in enumerate(team_users, start=1):
            if u.id == user_id:
                ranking["posicao"] = idx
                if idx > 1 and idx <= len(team_users):
                    ranking["xpParaProximo"] = team_users[idx - 2].totalXp - (user.totalXp or 0)

            if idx <= 3:
                ranking["top3"].append({
                    "nome": u.nome,
                    "avatarUrl": u.avatarUrl,
                    "totalXp": u.totalXp or 0,
                    "nivel": u.nivel or 1,
                    "posicao": idx
                })

        if len(team_users) > 0:
            ranking["xpDoLider"] = team_users[0].totalXp or 0

    # Último check-in bio
    all_checkins_user = await prisma.checkinbio.find_many(
        where={"idUsuario": user_id}
    )
    # Sort by date descending and get first
    ultimo_checkin = None
    if all_checkins_user:
        sorted_checkins = sorted(all_checkins_user, key=lambda c: c.dataHora if c.dataHora else datetime.min, reverse=True)
        ultimo_checkin = sorted_checkins[0] if sorted_checkins else None

    checkin_bio = None
    if ultimo_checkin:
        checkin_bio = {
            "nivelFoco": ultimo_checkin.nivelFoco or 0,
            "nivelEstresse": ultimo_checkin.nivelEstresse or 0,
            "horasSono": ultimo_checkin.horasSono or 0,
            "qualidadeSono": ultimo_checkin.qualidadeSono or 0,
            "dataHora": _iso(ultimo_checkin.dataHora)
        }

    # Cursos ativos (em andamento ou atrasados)
    cursos_ativos = []
    for m in matriculas:
        if m.status in ["EM_ANDAMENTO", "ATRASADO"]:
            cursos_ativos.append({
                "id": m.id,
                "cursoId": m.idCurso,
                "titulo": m.curso.titulo if m.curso else "Sem título",
                "thumbnailUrl": m.curso.thumbnailUrl if m.curso else None,
                "progresso": m.progresso or 0,
                "status": m.status,
                "prazo": _iso(m.prazo),
                "ultimoAcesso": _iso(m.ultimoAcesso),
                "ehObrigatorio": m.ehObrigatorio or False
            })

    # Ordenar: atrasados primeiro, depois por progresso decrescente
    cursos_ativos.sort(key=lambda x: (x["status"] != "ATRASADO", -x["progresso"]))

    # Cursos recomendados (cursos que o usuário não está matriculado)
    cursos_matriculados_ids = [m.idCurso for m in matriculas]
    cursos_recomendados = await prisma.materialfonte.find_many(
        where={"id": {"notIn": cursos_matriculados_ids}} if cursos_matriculados_ids else {},
        take=6
    )

    cursos_rec_list = []
    for c in cursos_recomendados:
        cursos_rec_list.append({
            "id": c.id,
            "titulo": c.titulo,
            "descricao": c.descricao,
            "categoria": c.categoria,
            "thumbnailUrl": c.thumbnailUrl,
            "tipoCurso": c.tipoCurso
        })

    return {
        "stats": stats,
        "ranking": ranking,
        "checkinBio": checkin_bio,
        "cursosAtivos": cursos_ativos[:6],  # Limitar a 6
        "cursosRecomendados": cursos_rec_list
    }


@app.get("/api/dashboard/manager")
async def dashboard_manager():
    '''
    Retorna dados agregados para o dashboard do gestor
    '''
    # KPIs principais
    total_colaboradores = await prisma.usuario.count(
        where={"papel": "COLABORADOR"}
    )

    total_matriculas = await prisma.matricula.count()
    matriculas_concluidas = await prisma.matricula.count(
        where={"status": "CONCLUIDO"}
    )

    taxa_conclusao = 0
    if total_matriculas > 0:
        taxa_conclusao = int((matriculas_concluidas / total_matriculas) * 100)

    cursos_atrasados = await prisma.matricula.count(
        where={"status": "ATRASADO"}
    )

    # Média de progresso (apenas cursos em andamento e atrasados)
    matriculas_ativas = await prisma.matricula.find_many(
        where={"status": {"in": ["EM_ANDAMENTO", "ATRASADO"]}}
    )

    media_progresso = 0
    if len(matriculas_ativas) > 0:
        media_progresso = int(
            sum(m.progresso or 0 for m in matriculas_ativas) / len(matriculas_ativas)
        )

    kpis = {
        "totalColaboradores": total_colaboradores,
        "taxaConclusao": taxa_conclusao,
        "cursosAtrasados": cursos_atrasados,
        "mediaProgresso": media_progresso
    }

    # Distribuição de status
    status_counts = {}
    all_matriculas = await prisma.matricula.find_many()

    for m in all_matriculas:
        status = m.status or "NAO_INICIADO"
        status_counts[status] = status_counts.get(status, 0) + 1

    distribuicao_status = {
        "NAO_INICIADO": status_counts.get("NAO_INICIADO", 0),
        "EM_ANDAMENTO": status_counts.get("EM_ANDAMENTO", 0),
        "CONCLUIDO": status_counts.get("CONCLUIDO", 0),
        "ATRASADO": status_counts.get("ATRASADO", 0),
        "REPROVADO_SIMULADO": status_counts.get("REPROVADO_SIMULADO", 0)
    }

    # Alertas críticos

    # 1. Colaboradores com cursos atrasados
    colaboradores_atrasados_data = await prisma.matricula.find_many(
        where={"status": "ATRASADO"},
        include={"usuario": True}
    )

    # Agrupar por usuário
    user_delays = {}
    for m in colaboradores_atrasados_data:
        if m.usuario:
            user_id = m.usuario.id
            if user_id not in user_delays:
                user_delays[user_id] = {
                    "id": m.usuario.id,
                    "nome": m.usuario.nome,
                    "avatarUrl": m.usuario.avatarUrl,
                    "cursosAtrasados": 0
                }
            user_delays[user_id]["cursosAtrasados"] += 1

    colaboradores_atrasados = sorted(
        user_delays.values(),
        key=lambda x: x["cursosAtrasados"],
        reverse=True
    )[:5]

    # 2. Cursos próximos do prazo (< 7 dias)
    prazo_proximo = datetime.utcnow() + timedelta(days=7)
    matriculas_prazo_proximo_all = await prisma.matricula.find_many(
        where={
            "status": {"in": ["NAO_INICIADO", "EM_ANDAMENTO"]},
            "prazo": {"lte": prazo_proximo, "gte": datetime.utcnow()}
        },
        include={"usuario": True, "curso": True}
    )
    # Sort by prazo in Python and take top 5
    matriculas_prazo_proximo = sorted(
        matriculas_prazo_proximo_all,
        key=lambda m: m.prazo if m.prazo else datetime.max
    )[:5]

    prazos_proximos = []
    for m in matriculas_prazo_proximo:
        if m.prazo:
            dias_restantes = (m.prazo - datetime.utcnow()).days
            prazos_proximos.append({
                "nomeUsuario": m.usuario.nome if m.usuario else "Desconhecido",
                "tituloCurso": m.curso.titulo if m.curso else "Sem título",
                "prazo": _iso(m.prazo),
                "diasRestantes": dias_restantes
            })

    # 3. Reprovados em simulados
    reprovados = await prisma.matricula.find_many(
        where={"status": "REPROVADO_SIMULADO"},
        include={"usuario": True, "curso": True},
        take=5
    )

    reprovados_list = []
    for m in reprovados:
        reprovados_list.append({
            "nomeUsuario": m.usuario.nome if m.usuario else "Desconhecido",
            "tituloCurso": m.curso.titulo if m.curso else "Sem título",
            "notaFinal": m.notaFinal
        })

    alertas = {
        "atrasados": colaboradores_atrasados,
        "prazosProximos": prazos_proximos,
        "reprovados": reprovados_list
    }

    # Performance: Top 5 e Bottom 5
    all_colaboradores = await prisma.usuario.find_many(
        where={"papel": "COLABORADOR"}
    )
    # Sort by XP
    sorted_by_xp = sorted(all_colaboradores, key=lambda u: u.totalXp or 0, reverse=True)
    top_performers = sorted_by_xp[:5]
    bottom_performers = sorted_by_xp[-5:]

    performance = {
        "top5": [
            {
                "nome": u.nome,
                "totalXp": u.totalXp or 0,
                "nivel": u.nivel or 1,
                "avatarUrl": u.avatarUrl
            }
            for u in top_performers
        ],
        "bottom5": [
            {
                "nome": u.nome,
                "totalXp": u.totalXp or 0,
                "nivel": u.nivel or 1,
                "avatarUrl": u.avatarUrl
            }
            for u in bottom_performers
        ]
    }

    # Bem-estar da equipe (última leitura de cada colaborador)
    all_checkins_raw = await prisma.checkinbio.find_many(
        include={"usuario": True}
    )
    # Sort by date descending
    all_checkins = sorted(all_checkins_raw, key=lambda c: c.dataHora if c.dataHora else datetime.min, reverse=True)

    # Filtrar último checkin de cada colaborador
    user_latest_checkin = {}
    for cb in all_checkins:
        if cb.usuario and cb.usuario.papel == "COLABORADOR":
            user_id = cb.usuario.id
            if user_id not in user_latest_checkin:
                user_latest_checkin[user_id] = cb

    checkins_recentes = list(user_latest_checkin.values())

    foco_medio = 0
    stress_medio = 0
    if len(checkins_recentes) > 0:
        foco_medio = int(
            sum(c.nivelFoco or 0 for c in checkins_recentes) / len(checkins_recentes)
        )
        stress_medio = int(
            sum(c.nivelEstresse or 0 for c in checkins_recentes) / len(checkins_recentes)
        )

    bem_estar = {
        "focoMedio": foco_medio,
        "stressMedio": stress_medio
    }

    # Comparativo de equipes
    equipes = await prisma.equipe.find_many(
        include={"usuarios": {"include": {"matriculas": True}}}
    )

    equipes_stats = []
    for e in equipes:
        usuarios_equipe = e.usuarios or []
        total_colab = len(usuarios_equipe)

        if total_colab == 0:
            continue

        xp_medio = sum(u.totalXp or 0 for u in usuarios_equipe) / total_colab

        # Calcular taxa de conclusão da equipe
        total_mat = sum(len(u.matriculas or []) for u in usuarios_equipe)
        concluidas_mat = sum(
            len([m for m in (u.matriculas or []) if m.status == "CONCLUIDO"])
            for u in usuarios_equipe
        )
        taxa_equipe = int((concluidas_mat / total_mat) * 100) if total_mat > 0 else 0

        # Cursos atrasados da equipe
        atrasados_equipe = sum(
            len([m for m in (u.matriculas or []) if m.status == "ATRASADO"])
            for u in usuarios_equipe
        )

        equipes_stats.append({
            "nome": e.nome,
            "totalColaboradores": total_colab,
            "taxaConclusao": taxa_equipe,
            "cursosAtrasados": atrasados_equipe,
            "xpMedio": int(xp_medio)
        })

    # Timeline de atividades recentes (logs de auditoria)
    logs_todos = await prisma.logauditoria.find_many(
        include={"usuario": True}
    )
    # Sort by date descending and take 10
    logs_recentes = sorted(logs_todos, key=lambda log: log.dataHora if log.dataHora else datetime.min, reverse=True)[:10]

    timeline = []
    for log in logs_recentes:
        timeline.append({
            "acao": log.acao,
            "detalhes": log.detalhes,
            "dataHora": _iso(log.dataHora),
            "usuarioNome": log.usuario.nome if log.usuario else "Sistema"
        })

    neuro_predictor = await predictive_lab.organization_snapshot(prisma)
    social_impact = await build_social_impact(prisma)

    return {
        "kpis": kpis,
        "distribuicaoStatus": distribuicao_status,
        "alertas": alertas,
        "performance": performance,
        "bemEstar": bem_estar,
        "equipes": equipes_stats,
        "timeline": timeline,
        "neuroPredictor": neuro_predictor,
        "socialImpact": social_impact,
    }



class CheckInPayload(BaseModel):
    userId: str
    horasSono: Optional[float] = None
    qualidadeSono: Optional[int] = None
    nivelFoco: int
    nivelEstresse: int
    nivelFadiga: Optional[int] = None
    origemDados: str = "MANUAL"
    dadosBrutosSensor: Optional[str] = None


@app.post("/api/checkins", status_code=201)
async def create_checkin(payload: CheckInPayload):
    checkin_id = _uuid()
    now = datetime.utcnow()

    # Calculate day of week (0=Monday, 6=Sunday) and hour
    dia_da_semana = now.weekday()
    hora_do_dia = now.hour

    await prisma.checkinbio.create(
        data={
            "id": checkin_id,
            "idUsuario": payload.userId,
            "horasSono": payload.horasSono,
            "qualidadeSono": payload.qualidadeSono,
            "nivelFoco": payload.nivelFoco,
            "nivelEstresse": payload.nivelEstresse,
            "nivelFadiga": payload.nivelFadiga,
            "origemDados": payload.origemDados,
            "dadosBrutosSensor": payload.dadosBrutosSensor,
            "dataHora": now,
            "diaDaSemana": dia_da_semana,
            "horaDoDia": hora_do_dia
        }
    )

    return {"id": checkin_id}


@app.get("/api/users/{user_id}/bio-analytics")
async def get_user_bio_analytics(user_id: str):
    """Analytics completo dos checkins biométricos de um usuário"""
    now = datetime.utcnow()
    all_checkins = await prisma.checkinbio.find_many(where={"idUsuario": user_id}, order={"dataHora": "desc"})
    if not all_checkins:
        return {"hasData": False, "lastCheckin": None, "weeklyAvg": None, "trends": None, "healthScore": None, "alerts": [], "patterns": None, "correlations": None}
    last_checkin = all_checkins[0]
    last_checkin_data = {"nivelFoco": last_checkin.nivelFoco, "nivelEstresse": last_checkin.nivelEstresse, "horasSono": last_checkin.horasSono, "qualidadeSono": last_checkin.qualidadeSono, "dataHora": _iso(last_checkin.dataHora) if last_checkin.dataHora else None}
    seven_days_ago = now - timedelta(days=7)
    checkins_7d = [c for c in all_checkins if c.dataHora and c.dataHora >= seven_days_ago]
    fourteen_days_ago = now - timedelta(days=14)
    checkins_14d = [c for c in all_checkins if c.dataHora and c.dataHora >= fourteen_days_ago]
    checkins_prev_week = [c for c in checkins_14d if c.dataHora and c.dataHora < seven_days_ago]
    weekly_avg = None
    if checkins_7d:
        avg_foco = sum(c.nivelFoco or 0 for c in checkins_7d) / len(checkins_7d)
        avg_stress = sum(c.nivelEstresse or 0 for c in checkins_7d) / len(checkins_7d)
        checkins_with_sleep = [c for c in checkins_7d if c.horasSono is not None]
        avg_sono = sum(c.horasSono for c in checkins_with_sleep) / len(checkins_with_sleep) if checkins_with_sleep else None
        checkins_with_quality = [c for c in checkins_7d if c.qualidadeSono is not None]
        avg_quality = sum(c.qualidadeSono for c in checkins_with_quality) / len(checkins_with_quality) if checkins_with_quality else None
        weekly_avg = {"foco": round(avg_foco, 1), "estresse": round(avg_stress, 1), "horasSono": round(avg_sono, 1) if avg_sono else None, "qualidadeSono": round(avg_quality, 1) if avg_quality else None}
    trends = None
    if checkins_7d and checkins_prev_week:
        avg_foco_current = sum(c.nivelFoco or 0 for c in checkins_7d) / len(checkins_7d)
        avg_foco_prev = sum(c.nivelFoco or 0 for c in checkins_prev_week) / len(checkins_prev_week)
        foco_diff = avg_foco_current - avg_foco_prev
        avg_stress_current = sum(c.nivelEstresse or 0 for c in checkins_7d) / len(checkins_7d)
        avg_stress_prev = sum(c.nivelEstresse or 0 for c in checkins_prev_week) / len(checkins_prev_week)
        stress_diff = avg_stress_current - avg_stress_prev
        def get_trend(diff):
            return "stable" if abs(diff) < 5 else ("up" if diff > 0 else "down")
        trends = {"foco": {"direction": get_trend(foco_diff), "percentage": round(foco_diff, 1)}, "estresse": {"direction": get_trend(stress_diff), "percentage": round(stress_diff, 1)}}
    health_score = None
    if weekly_avg:
        foco_component = weekly_avg["foco"] * 0.4
        stress_component = (100 - weekly_avg["estresse"]) * 0.4
        sleep_component = (weekly_avg["qualidadeSono"] * 10 * 0.2) if weekly_avg["qualidadeSono"] else 0
        score = foco_component + stress_component + sleep_component
        health_score = round(score, 0)
    alerts = []
    if len(checkins_7d) >= 3:
        recent_3 = checkins_7d[:3]
        if all(c.nivelEstresse and c.nivelEstresse > 70 for c in recent_3):
            alerts.append({"type": "high_stress", "severity": "critical", "message": "Estresse alto por 3+ dias consecutivos", "value": round(sum(c.nivelEstresse for c in recent_3) / 3, 1)})
    if len(checkins_7d) >= 5:
        recent_5 = checkins_7d[:5]
        if all(c.nivelFoco and c.nivelFoco < 50 for c in recent_5):
            alerts.append({"type": "low_focus", "severity": "warning", "message": "Foco abaixo de 50% por 5+ dias", "value": round(sum(c.nivelFoco for c in recent_5) / 5, 1)})
    if last_checkin.dataHora:
        days_since_last = (now - last_checkin.dataHora).days
        if days_since_last >= 3:
            alerts.append({"type": "no_checkin", "severity": "info", "message": f"Sem checkins há {days_since_last} dias", "value": days_since_last})
    if weekly_avg and weekly_avg["horasSono"] and weekly_avg["horasSono"] < 6:
        alerts.append({"type": "insufficient_sleep", "severity": "warning", "message": "Média de sono abaixo de 6h", "value": weekly_avg["horasSono"]})
    patterns = None
    if len(checkins_7d) >= 3:
        day_focus = {}
        for c in checkins_7d:
            day = c.diaDaSemana
            if day not in day_focus:
                day_focus[day] = []
            day_focus[day].append(c.nivelFoco or 0)
        day_avg = {day: sum(focus_list) / len(focus_list) for day, focus_list in day_focus.items()}
        best_day = max(day_avg.items(), key=lambda x: x[1]) if day_avg else None
        hour_focus = {}
        for c in checkins_7d:
            hour = c.horaDoDia
            if hour not in hour_focus:
                hour_focus[hour] = []
            hour_focus[hour].append(c.nivelFoco or 0)
        hour_avg = {hour: sum(focus_list) / len(focus_list) for hour, focus_list in hour_focus.items()}
        best_hour = max(hour_avg.items(), key=lambda x: x[1]) if hour_avg else None
        day_names = ["Segunda", "Terça", "Quarta", "Quinta", "Sexta", "Sábado", "Domingo"]
        patterns = {"bestDay": {"day": day_names[best_day[0]] if best_day else None, "avgFocus": round(best_day[1], 1) if best_day else None}, "bestHour": {"hour": best_hour[0] if best_hour else None, "avgFocus": round(best_hour[1], 1) if best_hour else None}}
    correlations = None
    user = await prisma.usuario.find_unique(where={"id": user_id}, include={"matriculas": True})
    if user and user.matriculas:
        total_courses = len(user.matriculas)
        completed = len([m for m in user.matriculas if m.status == "CONCLUIDO"])
        completion_rate = (completed / total_courses * 100) if total_courses > 0 else 0
        if weekly_avg:
            correlations = {"completionRate": round(completion_rate, 1), "avgFocus": weekly_avg["foco"], "interpretation": None}
            if weekly_avg["foco"] >= 70 and completion_rate >= 70:
                correlations["interpretation"] = "Alto foco correlacionado com boa performance"
            elif weekly_avg["foco"] < 50 and completion_rate < 50:
                correlations["interpretation"] = "Baixo foco pode estar impactando performance"
            elif weekly_avg["foco"] >= 70 and completion_rate < 50:
                correlations["interpretation"] = "Bom foco, mas performance abaixo do esperado"
            else:
                correlations["interpretation"] = "Performance independente do foco atual"
    return {"hasData": True, "lastCheckin": last_checkin_data, "weeklyAvg": weekly_avg, "trends": trends, "healthScore": health_score, "alerts": alerts, "patterns": patterns, "correlations": correlations}


if __name__ == "__main__":
    import uvicorn

    uvicorn.run(app, host="0.0.0.0", port=8000)
